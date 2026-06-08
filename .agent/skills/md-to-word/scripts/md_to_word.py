#!/usr/bin/env python3
"""
MD to Word Converter
Convert Markdown files to Word documents (.docx), preserving formatting

Usage:
    python md_to_word.py input.md
    python md_to_word.py input.md --output output.docx
    python md_to_word.py *.md  # Batch conversion
"""

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from typing import List, Optional


def parse_markdown_to_elements(md_content: str) -> List[dict]:
    """
    Parse Markdown content into a structured list of elements

    Args:
        md_content: Markdown text content

    Returns:
        List of elements, each containing type and content
    """
    elements = []
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings (# H1, ## H2, ### H3, etc.)
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line[level:].strip()
            elements.append({
                'type': 'heading',
                'level': level,
                'text': text
            })
            i += 1
            continue

        # Code blocks (```)
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip() or 'text'
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({
                'type': 'code',
                'lang': lang,
                'content': '\n'.join(code_lines)
            })
            i += 1
            continue

        # Horizontal rules (--- or ***)
        if line.strip() in ['---', '***', '___']:
            elements.append({'type': 'hr'})
            i += 1
            continue

        # Block quotes (>)
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            elements.append({
                'type': 'quote',
                'text': quote_text
            })
            i += 1
            continue

        # Unordered lists (- or *)
        if re.match(r'^\s*[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s+', lines[i]):
                item_text = re.sub(r'^\s*[-*]\s+', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            elements.append({
                'type': 'unordered_list',
                'items': items
            })
            continue

        # Ordered lists (1., 2., etc.)
        if re.match(r'^\s*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                item_text = re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            elements.append({
                'type': 'ordered_list',
                'items': items
            })
            continue

        # Tables (Markdown tables start with |)
        if line.strip().startswith('|') and '|' in line.strip():
            # Parse table
            rows = []
            # Header row
            header_row = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            rows.append(header_row)
            i += 1

            # Separator row (contains ---)
            if i < len(lines) and re.match(r'^\s*\|[\s\|:-]+\|\s*$', lines[i]):
                i += 1

            # Table content rows
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [cell.strip() for cell in lines[i].strip().split('|')[1:-1]]
                if cells:  # Avoid empty rows
                    rows.append(cells)
                i += 1

            if len(rows) > 1:  # At least header and one content row
                elements.append({
                    'type': 'table',
                    'rows': rows
                })
            continue

        # Regular paragraphs
        paragraph_lines = []
        while i < len(lines) and lines[i].strip():
            # Check if it's another special element
            if (lines[i].startswith('#') or
                lines[i].strip().startswith('```') or
                lines[i].strip() in ['---', '***', '___'] or
                lines[i].strip().startswith('>') or
                lines[i].strip().startswith('|') or
                re.match(r'^\s*[-*]\s+', lines[i]) or
                re.match(r'^\s*\d+\.\s+', lines[i])):
                break
            paragraph_lines.append(lines[i])
            i += 1

        if paragraph_lines:
            paragraph_text = ' '.join(paragraph_lines).strip()
            # Don't process inline formatting at parsing stage, preserve raw text
            elements.append({
                'type': 'paragraph',
                'text': paragraph_text
            })

    return elements


def process_inline_formatting(text: str) -> list:
    """
    Process inline formatting (bold, italic, code), return Run object configuration list

    Args:
        text: Raw text containing Markdown format markers

    Returns:
        List of Run configurations, each containing (text, bold, italic, code) attributes
    """
    runs = []
    remaining = text

    # Regex patterns for various inline formats
    # Note: Bold ** must be matched before italic * to avoid conflicts
    patterns = [
        (r'\*\*(.+?)\*\*', {'bold': True}),      # Bold **text**
        (r'\*(.+?)\*', {'italic': True}),        # Italic *text*
        (r'`(.+?)`', {'code': True}),            # Inline code `code`
    ]

    while remaining:
        # Find the earliest matched format
        earliest_match = None
        earliest_pos = len(remaining)
        match_type = None

        for pattern, style in patterns:
            match = re.search(pattern, remaining)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                match_type = style

        if earliest_match:
            # Add plain text before the match
            if earliest_pos > 0:
                runs.append({
                    'text': remaining[:earliest_pos],
                    'bold': False,
                    'italic': False,
                    'code': False
                })

            # Add formatted text
            runs.append({
                'text': earliest_match.group(1),
                'bold': match_type.get('bold', False),
                'italic': match_type.get('italic', False),
                'code': match_type.get('code', False)
            })

            # Continue processing remaining text
            remaining = remaining[earliest_match.end():]
        else:
            # No more format markers, add remaining text
            if remaining:
                runs.append({
                    'text': remaining,
                    'bold': False,
                    'italic': False,
                    'code': False
                })
            break

    return runs


def add_formatted_paragraph(doc, text: str, base_size: int = 11):
    """
    Add paragraph with inline formatting to Word document

    Args:
        doc: Word document object
        text: Text containing Markdown formatting
        base_size: Base font size
    """
    from docx.shared import Pt

    p = doc.add_paragraph()
    runs_config = process_inline_formatting(text)

    for run_config in runs_config:
        run = p.add_run(run_config['text'])

        # Set font size
        run.font.size = Pt(base_size)

        # Apply formatting
        if run_config['bold']:
            run.font.bold = True
        if run_config['italic']:
            run.font.italic = True
        if run_config['code']:
            run.font.name = 'Courier New'
            run.font.size = Pt(base_size - 1)  # Code font slightly smaller

    return p


def setup_logging() -> logging.Logger:
    """
    Set up logger, log file saved in the same directory as the Python script

    Returns:
        Configured logger
    """
    # Get script directory
    script_dir = Path(__file__).parent
    log_path = script_dir / 'md_to_word.log'

    # Create logger
    logger = logging.getLogger('md_to_word')
    logger.setLevel(logging.DEBUG)

    # Clear existing handlers
    logger.handlers.clear()

    # File handler
    file_handler = logging.FileHandler(log_path, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(file_formatter)

    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s: %(message)s')
    console_handler.setFormatter(console_formatter)

    # Add handlers
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

    logger.debug(f'Log file: {log_path}')
    return logger


def convert_to_word(md_path: str, output_path: Optional[str] = None) -> str:
    """
    Convert Markdown file to Word document

    Args:
        md_path: Markdown file path
        output_path: Output Word file path (optional, defaults to original filename)

    Returns:
        Generated Word file path
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print("Error: python-docx library missing")
        print("Please install: pip install python-docx")
        sys.exit(1)

    # Determine output path
    md_file = Path(md_path)
    if output_path is None:
        output_path = md_file.with_suffix('.docx')

    # Set up logging
    logger = setup_logging()
    logger.info(f'Starting conversion: {md_path}')
    logger.debug(f'Output path: {output_path}')

    # Read Markdown file
    logger.debug(f'Reading file: {md_path}')
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    logger.debug(f'File size: {len(md_content)} characters')

    # Parse Markdown
    logger.debug('Parsing Markdown content')
    elements = parse_markdown_to_elements(md_content)
    logger.debug(f'Parsing complete, {len(elements)} elements found')

    # Create Word document
    logger.debug('Creating Word document')
    doc = Document()

    # 添加元素到 Word 文档
    element_counts = {
        'heading': 0,
        'paragraph': 0,
        'unordered_list': 0,
        'ordered_list': 0,
        'code': 0,
        'quote': 0,
        'table': 0,
        'hr': 0
    }

    for element in elements:
        if element['type'] == 'heading':
            # Heading
            level = min(element['level'], 3)  # Word only supports H1-H3
            heading = doc.add_heading(element['text'], level=level)
            heading.style.font.size = Pt(16 - level * 2)
            heading.style.font.bold = True
            element_counts['heading'] += 1

        elif element['type'] == 'paragraph':
            # Regular paragraph (supports inline formatting)
            add_formatted_paragraph(doc, element['text'], base_size=11)
            element_counts['paragraph'] += 1

        elif element['type'] == 'unordered_list':
            # Unordered list (supports inline formatting)
            for item in element['items']:
                p = doc.add_paragraph(style='List Bullet')
                runs_config = process_inline_formatting(item)
                for run_config in runs_config:
                    run = p.add_run(run_config['text'])
                    run.font.size = Pt(11)
                    if run_config['bold']:
                        run.font.bold = True
                    if run_config['italic']:
                        run.font.italic = True
                    if run_config['code']:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            element_counts['unordered_list'] += 1

        elif element['type'] == 'ordered_list':
            # Ordered list (supports inline formatting)
            for item in element['items']:
                p = doc.add_paragraph(style='List Number')
                runs_config = process_inline_formatting(item)
                for run_config in runs_config:
                    run = p.add_run(run_config['text'])
                    run.font.size = Pt(11)
                    if run_config['bold']:
                        run.font.bold = True
                    if run_config['italic']:
                        run.font.italic = True
                    if run_config['code']:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            element_counts['ordered_list'] += 1

        elif element['type'] == 'code':
            # Code block
            p = doc.add_paragraph()
            run = p.add_run(f"[Code: {element['lang']}]\n{element['content']}")
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            element_counts['code'] += 1

        elif element['type'] == 'quote':
            # Quote (supports inline formatting)
            p = doc.add_paragraph()
            runs_config = process_inline_formatting(element['text'])
            for run_config in runs_config:
                run = p.add_run(run_config['text'])
                run.font.size = Pt(11)
                run.font.italic = True  # Quote uses italic
                if run_config['bold']:
                    run.font.bold = True
                if run_config['code']:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            p.paragraph_format.left_indent = Pt(20)
            element_counts['quote'] += 1

        elif element['type'] == 'table':
            # Table
            rows = element['rows']
            if not rows:
                continue

            # Create table
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'

            # Fill table content
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]

                    # Process inline formatting
                    runs_config = process_inline_formatting(cell_text)

                    # Clear default paragraph and add formatted content
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        p.clear()
                    else:
                        p = cell.add_paragraph()

                    for run_config in runs_config:
                        run = p.add_run(run_config['text'])
                        run.font.size = Pt(10)

                        # Bold header row
                        if row_idx == 0:
                            run.font.bold = True

                        # Apply inline formatting
                        if run_config['bold']:
                            run.font.bold = True
                        if run_config['italic']:
                            run.font.italic = True
                        if run_config['code']:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
            element_counts['table'] += 1

        elif element['type'] == 'hr':
            # Horizontal rule
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(12)
            element_counts['hr'] += 1

    logger.debug(f'Element counts: {element_counts}')

    # Save Word document
    logger.debug(f'Saving Word document: {output_path}')
    doc.save(output_path)
    logger.info(f'✓ Converted: {md_path} -> {output_path}')
    logger.debug('Conversion complete')

    return str(output_path)


def convert_to_word_with_error_handling(md_path: str, output_path: Optional[str] = None) -> str:
    """
    Conversion function with exception handling

    Args:
        md_path: Markdown file path
        output_path: Output Word file path

    Returns:
        Generated Word file path

    Raises:
        Exception: Thrown when conversion fails
    """
    try:
        return convert_to_word(md_path, output_path)
    except FileNotFoundError as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'File not found: {md_path} - {e}')
        raise
    except PermissionError as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'Permission error: {md_path} - {e}')
        raise
    except Exception as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'Conversion failed: {md_path} - {type(e).__name__}: {e}', exc_info=True)
        raise


def main():
    """Main function"""
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 md_to_word.py README.md
  python3 md_to_word.py README.md --output README_word.docx
  python3 md_to_word.py *.md              # Batch conversion
  python3 md_to_word.py docs/*.md         # Convert specified directory
        """
    )

    parser.add_argument(
        'input_files',
        nargs='+',
        help='Input Markdown file paths (supports wildcards)'
    )

    parser.add_argument(
        '--output', '-o',
        help='Output Word file path (only for single file)'
    )

    parser.add_argument(
        '--recursive', '-r',
        action='store_true',
        help='Recursively process subdirectories'
    )

    args = parser.parse_args()

    # Set up global logging
    main_logger = logging.getLogger('md_to_word_main')
    main_logger.setLevel(logging.INFO)
    main_logger.handlers.clear()

    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s: %(message)s')
    console_handler.setFormatter(console_formatter)
    main_logger.addHandler(console_handler)

    main_logger.info('MD to Word Converter started')

    # Expand wildcards
    input_files = []
    for pattern in args.input_files:
        if '*' in pattern or '?' in pattern:
            matched = list(Path('.').glob(pattern))
            if args.recursive:
                matched.extend(Path('.').rglob(pattern.split('/')[-1]))
            input_files.extend(matched)
        else:
            input_files.append(Path(pattern))

    # Filter non-existent files
    valid_files = [f for f in input_files if f.exists() and f.suffix.lower() == '.md']

    if not valid_files:
        main_logger.error("Error: No valid Markdown files found")
        sys.exit(1)

    main_logger.info(f"Found {len(valid_files)} Markdown file(s)")

    # Convert files
    success_count = 0
    for md_file in valid_files:
        try:
            output_path = args.output if len(valid_files) == 1 and args.output else None
            convert_to_word(str(md_file), output_path)
            success_count += 1
        except Exception as e:
            main_logger.error(f"✗ Conversion failed: {md_file} - {e}")

    main_logger.info(f"Complete! Successfully converted {success_count}/{len(valid_files)} file(s)")


if __name__ == '__main__':
    main()